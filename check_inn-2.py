#!/usr/bin/python3
# -*- coding: utf-8 -*-

import sys
import os
import re
import math
from datetime import datetime, timedelta
# Import docx, xlrd, openpyxl are handled inside functions, as in original script.

def calculate_birth_date(inn_digits):
    """
    Обчислює дату народження з перших 5 цифр РНОКПП.
    """
    try:
        # Перші 5 цифр - це кількість днів, що пройшли з 31.12.1899
        days_since = int("".join(map(str, inn_digits[:5])))
        
        # Базова дата для розрахунку
        base_date = datetime(1899, 12, 31)
        
        # Обчислюємо дату народження
        birth_date = base_date + timedelta(days=days_since)
        
        return birth_date.strftime("%d.%m.%Y")
    except Exception:
        return "Невідома дата"
        

def check_rnokpp(text): # Змінено check_inn на check_rnokpp для ясності
    """
    Перевіряє коректність 10-значного РНОКПП (ІНН фізичної особи).
    Повертає список помилок, а також список коректних РНОКПП та їхню дату народження.
    """
    results = {"errors": [], "valid": []} # valid міститиме кортежі (РНОКПП, Дата)
    
    weights = [ -1, 5, 7, 9, 4, 6, 10, 5, 7]
    
    # Пошук 10-значных чисел
    matches = re.findall(r'\b\d{10}\b', str(text))

    for match in matches:
        digits = [int(char) for char in match]

        # Расчет контрольной суммы (РНОКПП)
        k1 = sum(x * y for x, y in zip(digits[:9], weights))
        k2 = k1 % 11
        
        checksum = k2
        if k2 == 10:
            checksum = 0 # Якщо остача 10, контрольний розряд = 0

        # Сравнение контрольной суммы с последней цифрой
        if checksum == digits[9]:
            # Обчислюємо дату народження тільки для вірного коду
            birth_date = calculate_birth_date(digits) 
            results["valid"].append((match, birth_date)) 
        else:
            results["errors"].append(match)

    return results

def process_odt(file_path):
    from odf.opendocument import load
    from odf.text import P
    from odf.table import Table, TableRow, TableCell
    total_checked = 0
    total_errors = 0

    try:
        doc = load(file_path)
    except Exception as e:
        print(f"Помилка при відкритті файла ODT/ODS: {e}")
        return
    
    for elem in doc.getElementsByType(P):
        result = check_rnokpp(elem) 
        total_checked += len(result["valid"]) + len(result["errors"])
        total_errors += len(result["errors"])
        for error in result["errors"]:
            print(f"==> ERROR: {error}")
        for valid, date in result["valid"]: 
            print(f"ok: {valid} (Дата нар.: {date})") 

    print(f"Усього РНОКПП перевірено: {total_checked}, з помилками: {total_errors}")

def process_docx(file_path):
    import docx
    total_checked = 0
    total_errors = 0

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Помилка при відкритті файла DOCX: {e}")
        return

    # Перевірка параграфів
    for paragraph in doc.paragraphs:
        result = check_rnokpp(paragraph.text) 
        total_checked += len(result["valid"]) + len(result["errors"])
        total_errors += len(result["errors"])
        for error in result["errors"]:
            print(f"==> ERROR: {error}")
        for valid, date in result["valid"]: 
            print(f"ok: {valid} (Дата нар.: {date})") 

    # Перевірка таблиць
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                result = check_rnokpp(cell.text) 
                total_checked += len(result["valid"]) + len(result["errors"])
                total_errors += len(result["errors"])
                for error in result["errors"]:
                    print(f"==> ERROR: {error}")
                for valid, date in result["valid"]:
                    print(f"ok: {valid} (Дата нар.: {date})")

    print(f"Усього РНОКПП перевірено: {total_checked}, з помилками: {total_errors}")


def process_xls(file_path):
    import xlrd
    try:
        workbook = xlrd.open_workbook(file_path)
    except Exception as e:
        print(f"Помилка при відкритті файла XLS: {e}")
        return

    sheet = workbook.sheet_by_index(0)
    total_checked = 0
    total_errors = 0

    for row in range(sheet.nrows):
        for col in range(sheet.ncols):
            cell_value = sheet.cell_value(row, col)
            
            # Обробка числових значень
            if isinstance(cell_value, float) and cell_value == int(cell_value):
                cell_value = str(int(cell_value))
            elif isinstance(cell_value, float):
                cell_value = str(cell_value)
            
            if not cell_value:
                continue
                
            result = check_rnokpp(cell_value) # Змінено check_inn на check_rnokpp
            total_checked += len(result["valid"]) + len(result["errors"])
            total_errors += len(result["errors"])
            for error in result["errors"]:
                # Залишено "Невірний РНОКПП" для кращого пояснення помилки
                print(f"==> ERROR: {error} (Невірний РНОКПП)") 
            # ВИПРАВЛЕНО: Вивід коректний
            for valid, date in result["valid"]:
                print(f"ok: {valid} (Дата нар.: {date})")

    print(f"Усього РНОКПП перевірено: {total_checked}, Помилок: {total_errors}")

def process_xlsx(file_path):
    import openpyxl
    try:
        workbook = openpyxl.load_workbook(file_path)
    except Exception as e:
        print(f"Помилка при відкритті файла XLSX: {e}")
        return

    sheet = workbook.active
    total_checked = 0
    total_errors = 0

    for row in sheet.iter_rows():
        for cell in row:
            cell_value = cell.value
            
            # Обробка числових значень
            if isinstance(cell_value, int):
                cell_value = str(cell_value)
            elif isinstance(cell_value, float):
                if cell_value == int(cell_value):
                    cell_value = str(int(cell_value))
                else:
                    cell_value = str(cell_value)

            if cell_value is None:
                continue
                
            result = check_rnokpp(cell_value) # Змінено check_inn на check_rnokpp
            total_checked += len(result["valid"]) + len(result["errors"])
            total_errors += len(result["errors"])
            for error in result["errors"]:
                print(f"==> ERROR: {error}")
            # ВИПРАВЛЕНО: Вивід коректний
            for valid, date in result["valid"]:
                print(f"ok: {valid} (Дата нар.: {date})")

    print(f"Усього РНОКПП перевірено: {total_checked}, Помилок: {total_errors}")

def main():
    if getattr(sys, 'frozen', False):
        program_path = sys.executable
    else:
        program_path = os.path.abspath(__file__)
	
    if len(sys.argv) != 2:
        print(f"Використання: {program_path} <файл>")
        sys.exit(1)

    file_path = sys.argv[1]
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".docx" or extension == ".doc":
        process_docx(file_path)
    elif extension == ".xls":
        process_xls(file_path)
    elif extension == ".odt" or extension == ".ods":
        process_odt(file_path)
    elif extension == ".xlsx":
        process_xlsx(file_path)
    else:
        print("Не той формат (тільки doc, docx, xls, xlsx).")
        sys.exit(1)

if __name__ == "__main__":
    main()
